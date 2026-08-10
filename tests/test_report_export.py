from __future__ import annotations

import inspect
import io
from datetime import UTC, datetime
from types import SimpleNamespace

from docx import Document
from pypdf import PdfReader

from src.services.report_export import (
    ReportExportContext,
    ReportStyleSettings,
    build_report_docx,
    build_report_pdf,
    clean_report_markdown_for_display,
    project_report_context,
)
import src.services.report_export as report_export
from src.ui.theme import apply_theme


SAMPLE_MARKDOWN = """# 中国分子诊断行业研究报告

## 执行摘要

中国分子诊断行业的增长需要同时核验需求、支付与技术商业化条件。

## 行业定义与研究口径

> 核心市场不包含第三方医学检验服务，所有市场规模应使用同一统计口径。

- 产品范围：核酸与分子检测相关仪器、试剂及配套耗材
- 地理范围：中国大陆
- 证据来源：[国家药监局](https://www.nmpa.gov.cn/)

## 市场现状与竞争格局

主要结论必须区分**事实**、分析师推断与预测。

## 未来趋势

- 基准情景：临床需求持续扩展
- 加速情景：创新检测获得支付支持
- 受阻情景：价格和合规压力抑制商业化

#### 基准情景

在基准情景下，临床需求、支付能力与商业化效率需要同步观察。

| 维度 | 判断 | 证据状态 |
|---|---|---|
| 临床需求 | 稳步扩展 | 已审核 |
| 支付环境 | 仍需监测 | 存在缺口 |
"""


def context() -> ReportExportContext:
    return ReportExportContext(
        title="中国分子诊断行业研究报告",
        markdown=SAMPLE_MARKDOWN,
        industry="分子诊断",
        region="中国",
        time_horizon="2026-2030",
        report_status="经人工审核的通用行业研究报告",
        generated_at=datetime(2026, 7, 27, tzinfo=UTC),
        sop_label="Trident 专业行业研究 SOP v1.0.0",
    )


def test_word_export_is_editable_and_contains_report_content() -> None:
    payload = build_report_docx(context())

    assert payload.startswith(b"PK")
    document = Document(io.BytesIO(payload))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "中国分子诊断行业研究报告" in text
    assert "Trident 专业行业研究 SOP v1.0.0" in text
    assert "市场现状与竞争格局" in text
    assert "####" not in text
    assert len(document.tables) == 1
    assert document.tables[0].cell(1, 0).text == "临床需求"
    hyperlink_targets = [
        relationship.target_ref
        for relationship in document.part.rels.values()
        if relationship.reltype.endswith("/hyperlink")
    ]
    assert "https://www.nmpa.gov.cn/" in hyperlink_targets


def test_pdf_export_is_paginated_and_has_metadata() -> None:
    payload = build_report_pdf(context())

    assert payload.startswith(b"%PDF")
    reader = PdfReader(io.BytesIO(payload))
    assert len(reader.pages) >= 2
    assert reader.metadata.title == "中国分子诊断行业研究报告"


def test_pdf_export_uses_cloud_safe_cjk_fallback_without_system_font(
    monkeypatch,
) -> None:
    monkeypatch.setattr(report_export, "PDF_CJK_SANS_FONT", "MissingIndustryReportFont")
    monkeypatch.setattr(report_export, "PDF_SANS_FONT_CANDIDATES", ())
    monkeypatch.delenv("INDUSTRY_REPORT_CJK_FONT", raising=False)

    payload = report_export.build_report_pdf(context())

    assert payload.startswith(b"%PDF")
    assert len(PdfReader(io.BytesIO(payload)).pages) >= 2


def test_project_context_uses_the_projects_sop_version() -> None:
    project = SimpleNamespace(
        industry="工业机器人",
        region="全球",
        time_horizon="2026-2030",
        research_brief_artifact=SimpleNamespace(
            methodology=SimpleNamespace(
                sop_name="Trident 专业行业研究 SOP",
                sop_version="1.0.0",
            )
        ),
    )

    export = project_report_context(
        project,
        title="全球工业机器人行业研究",
        markdown="# 报告",
        report_status="经人工审核",
        generated_at=datetime(2026, 7, 27, tzinfo=UTC),
    )

    assert export.sop_label == "Trident 专业行业研究 SOP v1.0.0"


def test_download_button_theme_forces_white_text() -> None:
    source = inspect.getsource(apply_theme)

    assert ".stDownloadButton > button" in source
    assert "color: #FFFFFF !important" in source


def test_style_settings_are_applied_to_word_and_pdf_exports() -> None:
    styled = ReportStyleSettings(
        font_label="报告宋体",
        heading_color="#243B53",
        body_color="#52606D",
        report_title_size=38,
        level_one_size=28,
        level_two_size=21,
        body_size=17,
        line_height=1.9,
    )
    customized = ReportExportContext(**{**context().__dict__, "style": styled})

    word_payload = build_report_docx(customized)
    document = Document(io.BytesIO(word_payload))
    # OOXML stores type in half-points, so 12.75 pt is rounded to 12.5 pt.
    assert document.styles["Normal"].font.size.pt == 12.5
    assert str(document.styles["Normal"].font.color.rgb) == "52606D"
    assert document.styles["Heading 1"].font.size.pt == 21
    assert str(document.styles["Heading 1"].font.color.rgb) == "243B53"
    assert build_report_pdf(customized).startswith(b"%PDF")


def test_paragraphs_never_start_with_closing_punctuation_in_any_delivery() -> None:
    dirty = """# 标题

## 章节

，第一段结论。

。第二段结论。

…第三段结论。

—第四段结论。

- ；列表结论
"""
    cleaned = clean_report_markdown_for_display(dirty)
    visible_lines = [line.lstrip("#-* >") for line in cleaned.splitlines() if line.strip()]
    assert all(not line.startswith(tuple("，。；：！？、…—,.!?;:）》】〉」』’”)]")) for line in visible_lines)

    customized = ReportExportContext(**{**context().__dict__, "markdown": dirty})
    document = Document(io.BytesIO(build_report_docx(customized)))
    body_paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    assert all(not text.startswith(tuple("，。；：！？、…—,.!?;:）》】〉」』’”)]")) for text in body_paragraphs)
    assert build_report_pdf(customized).startswith(b"%PDF")
