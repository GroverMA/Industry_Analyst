"""Professional Word and PDF exports for human-reviewed research reports."""

from __future__ import annotations

import html
import io
import os
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ACCENT = "356B77"
ACCENT_DARK = "234D57"
INK = "172033"
MUTED = "667085"
LIGHT_FILL = "EEF5F5"
REPORT_FONT = "Calibri"
PDF_CJK_SANS_FONT = "IndustryReportCJKSans"
PDF_CJK_SERIF_FONT = "IndustryReportCJKSerif"
WORD_CJK_SANS_FONT_CANDIDATES = (
    ("Noto Sans CJK SC", "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
    ("Noto Sans CJK SC", "/usr/share/fonts/opentype/noto/NotoSansCJKsc-Regular.otf"),
    ("Arial Unicode MS", "/System/Library/Fonts/Supplemental/Arial Unicode.ttf"),
    ("Hiragino Sans GB", "/System/Library/Fonts/Hiragino Sans GB.ttc"),
)
WORD_CJK_SERIF_FONT_CANDIDATES = (
    ("Noto Serif CJK SC", "/usr/share/fonts/opentype/noto/NotoSerifCJK-Regular.ttc"),
    ("Noto Serif CJK SC", "/usr/share/fonts/opentype/noto/NotoSerifCJKsc-Regular.otf"),
    ("Songti SC", "/System/Library/Fonts/Supplemental/Songti.ttc"),
)
PDF_SANS_FONT_CANDIDATES = (
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/opentype/noto/NotoSansCJKsc-Regular.otf",
    "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
    "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
)
PDF_SERIF_FONT_CANDIDATES = (
    "/usr/share/fonts/opentype/noto/NotoSerifCJK-Regular.ttc",
    "/usr/share/fonts/opentype/noto/NotoSerifCJKsc-Regular.otf",
    "/System/Library/Fonts/Supplemental/Songti.ttc",
)

LEADING_PUNCTUATION = "，。；：！？、…—,.!?;:）)]】》〉」』’”"


def _resolve_word_cjk_font(
    candidates: tuple[tuple[str, str], ...],
    *,
    fallback: str,
) -> str:
    for family, location in candidates:
        if Path(location).is_file():
            return family
    return fallback


REPORT_CJK_SANS_FONT = _resolve_word_cjk_font(
    WORD_CJK_SANS_FONT_CANDIDATES,
    fallback="Microsoft YaHei",
)
REPORT_CJK_SERIF_FONT = _resolve_word_cjk_font(
    WORD_CJK_SERIF_FONT_CANDIDATES,
    fallback="SimSun",
)


def _hex(value: str, fallback: str) -> str:
    cleaned = str(value or "").strip().lstrip("#").upper()
    return cleaned if re.fullmatch(r"[0-9A-F]{6}", cleaned) else fallback


def _px_to_pt(value: int | float) -> float:
    return round(float(value) * 0.75, 2)


def _word_fonts(font_label: str) -> tuple[str, str]:
    if font_label == "报告宋体":
        return "Times New Roman", REPORT_CJK_SERIF_FONT
    if font_label == "系统字体":
        return "Arial", REPORT_CJK_SANS_FONT
    return "Arial", REPORT_CJK_SANS_FONT


def _strip_leading_punctuation(text: str) -> str:
    """Prevent closing punctuation from becoming the first visible glyph."""

    return re.sub(rf"^[\s{re.escape(LEADING_PUNCTUATION)}]+", "", str(text or ""))


def clean_report_markdown_for_display(markdown: str) -> str:
    """Normalize paragraph starts without damaging Markdown structure."""

    cleaned_lines: list[str] = []
    for raw_line in str(markdown or "").splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            cleaned_lines.append("")
            continue
        structural = re.match(r"^(\s*(?:#{1,6}\s+|[-*]\s+|>\s?))(.*)$", line)
        if structural:
            cleaned_lines.append(
                structural.group(1) + _strip_leading_punctuation(structural.group(2))
            )
            continue
        if line.lstrip().startswith("|"):
            cells = [
                cell if re.fullmatch(r":?-{3,}:?", cell.strip())
                else _strip_leading_punctuation(cell.strip())
                for cell in line.strip().strip("|").split("|")
            ]
            cleaned_lines.append("| " + " | ".join(cells) + " |")
            continue
        indent = line[: len(line) - len(line.lstrip())]
        cleaned_lines.append(indent + _strip_leading_punctuation(line.lstrip()))
    return "\n".join(cleaned_lines).strip()


@dataclass(frozen=True)
class ReportStyleSettings:
    """One typography contract shared by web preview, Word, and PDF."""

    font_label: str = "专业无衬线"
    heading_color: str = "#172033"
    body_color: str = "#3F4A5E"
    report_title_size: int = 34
    level_one_size: int = 26
    level_two_size: int = 20
    body_size: int = 16
    line_height: float = 1.85


@dataclass(frozen=True)
class ReportExportContext:
    title: str
    markdown: str
    industry: str
    region: str
    time_horizon: str
    report_status: str
    generated_at: datetime
    sop_label: str | None = None
    style: ReportStyleSettings = ReportStyleSettings()


def project_report_context(
    project,
    *,
    title: str,
    markdown: str,
    report_status: str,
    generated_at: datetime,
    style: ReportStyleSettings | None = None,
) -> ReportExportContext:
    """Create one export context for both quick and advanced workspaces."""

    methodology = getattr(project.research_brief_artifact, "methodology", None)
    sop_label = None
    if methodology is not None:
        sop_label = f"{methodology.sop_name} v{methodology.sop_version}"
    return ReportExportContext(
        title=title,
        markdown=markdown,
        industry=project.industry,
        region=project.region,
        time_horizon=project.time_horizon,
        report_status=report_status,
        generated_at=generated_at,
        sop_label=sop_label,
        style=style or ReportStyleSettings(),
    )


def build_report_docx(context: ReportExportContext) -> bytes:
    """Build a polished, editable Word report using business-report styling."""

    document = Document()
    _configure_document(document, context)
    _add_word_cover(document, context)
    document.add_page_break()
    _add_markdown_to_word(
        document,
        clean_report_markdown_for_display(context.markdown),
        context.style,
        skip_first_title=True,
    )

    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def build_report_pdf(context: ReportExportContext) -> bytes:
    """Build a paginated PDF without depending on LibreOffice at runtime."""

    pdf_font = _register_pdf_font(context.style.font_label)

    stream = io.BytesIO()
    document = SimpleDocTemplate(
        stream,
        pagesize=LETTER,
        rightMargin=0.82 * inch,
        leftMargin=0.82 * inch,
        topMargin=0.76 * inch,
        bottomMargin=0.72 * inch,
        title=context.title,
        author="Trident",
        subject=context.report_status,
    )
    styles = _pdf_styles(pdf_font, context.style)
    story = _pdf_cover(context, styles)
    story.append(PageBreak())
    story.extend(
        _markdown_to_pdf(
            clean_report_markdown_for_display(context.markdown),
            styles,
            skip_first_title=True,
        )
    )
    document.build(
        story,
        canvasmaker=lambda *args, **kwargs: _NumberedCanvas(
            *args,
            report_title=context.title,
            **kwargs,
        ),
    )
    return stream.getvalue()


def _register_pdf_font(font_label: str) -> str:
    """Embed a real CJK font so Chinese remains visible across PDF readers."""

    serif = font_label == "报告宋体"
    font_name = PDF_CJK_SERIF_FONT if serif else PDF_CJK_SANS_FONT

    try:
        pdfmetrics.getFont(font_name)
        return font_name
    except KeyError:
        pass

    configured = os.getenv("INDUSTRY_REPORT_CJK_FONT")
    candidates = ((configured,) if configured else ()) + (
        PDF_SERIF_FONT_CANDIDATES if serif else PDF_SANS_FONT_CANDIDATES
    )
    for candidate in candidates:
        font_path = Path(candidate).expanduser()
        if not font_path.is_file():
            continue
        try:
            pdfmetrics.registerFont(TTFont(font_name, str(font_path), subfontIndex=0))
        except Exception:
            continue
        return font_name

    # Streamlit Community Cloud does not guarantee a local CJK font file.
    # ReportLab's standard CID font keeps Chinese PDF generation available
    # without a system package or a private bundled asset.
    fallback_name = "STSong-Light"
    try:
        pdfmetrics.getFont(fallback_name)
    except KeyError:
        pdfmetrics.registerFont(UnicodeCIDFont(fallback_name))
    return fallback_name


def _configure_document(document: Document, context: ReportExportContext) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    document.core_properties.title = context.title
    document.core_properties.subject = context.report_status
    document.core_properties.author = "Trident"
    document.core_properties.keywords = "industry research, evidence, human review"

    styles = document.styles
    normal = styles["Normal"]
    latin_font, east_asia_font = _word_fonts(context.style.font_label)
    body_pt = _px_to_pt(context.style.body_size)
    heading_color = _hex(context.style.heading_color, INK)
    body_color = _hex(context.style.body_color, MUTED)
    _set_style_font(normal, latin_font, east_asia_font, body_pt, body_color)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(max(5, body_pt * 0.55))
    normal.paragraph_format.line_spacing = context.style.line_height

    heading_tokens = {
        "Heading 1": (_px_to_pt(context.style.level_one_size), heading_color, 18, 8),
        "Heading 2": (_px_to_pt(context.style.level_two_size), heading_color, 13, 6),
        "Heading 3": (_px_to_pt(max(17, context.style.level_two_size - 3)), heading_color, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        _set_style_font(style, latin_font, east_asia_font, size, color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    bullet = styles["List Bullet"]
    _set_style_font(bullet, latin_font, east_asia_font, body_pt, body_color)
    bullet.paragraph_format.left_indent = Inches(0.5)
    bullet.paragraph_format.first_line_indent = Inches(-0.25)
    bullet.paragraph_format.space_after = Pt(8)
    bullet.paragraph_format.line_spacing = context.style.line_height

    header = section.header
    header.is_linked_to_previous = False
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header_paragraph.add_run("TRIDENT  |  HUMAN-REVIEWED INDUSTRY RESEARCH")
    _set_run_font(run, latin_font, east_asia_font, 8.5, MUTED, bold=True)

    footer = section.footer
    footer.is_linked_to_previous = False
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    label = footer_paragraph.add_run("Trident  |  ")
    _set_run_font(label, latin_font, east_asia_font, 8.5, MUTED)
    _add_page_field(footer_paragraph)


def _add_word_cover(document: Document, context: ReportExportContext) -> None:
    latin_font, east_asia_font = _word_fonts(context.style.font_label)
    heading_color = _hex(context.style.heading_color, INK)
    body_color = _hex(context.style.body_color, MUTED)
    for _ in range(4):
        document.add_paragraph()
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("INDUSTRY RESEARCH REPORT")
    _set_run_font(run, latin_font, east_asia_font, 10, ACCENT, bold=True)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    title_run = title.add_run(context.title)
    _set_run_font(
        title_run,
        latin_font,
        east_asia_font,
        _px_to_pt(context.style.report_title_size),
        heading_color,
        bold=True,
    )

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(26)
    subtitle_run = subtitle.add_run(
        f"{context.industry}  |  {context.region}  |  {context.time_horizon}"
    )
    _set_run_font(subtitle_run, latin_font, east_asia_font, 13, ACCENT_DARK)

    status = document.add_paragraph()
    status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    status.paragraph_format.space_after = Pt(8)
    status_run = status.add_run(context.report_status)
    _set_run_font(status_run, latin_font, east_asia_font, 10.5, body_color, bold=True)

    if context.sop_label:
        sop = document.add_paragraph()
        sop.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sop.paragraph_format.space_after = Pt(4)
        sop_run = sop.add_run(f"Methodology: {context.sop_label}")
        _set_run_font(sop_run, latin_font, east_asia_font, 9.5, body_color)

    date = document.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date.add_run(context.generated_at.strftime("%Y-%m-%d"))
    _set_run_font(date_run, latin_font, east_asia_font, 10, body_color)


def _add_markdown_to_word(
    document: Document,
    markdown: str,
    style_settings: ReportStyleSettings,
    *,
    skip_first_title: bool,
) -> None:
    first_heading_skipped = False
    for kind, text, level in _markdown_blocks(markdown):
        if kind == "heading":
            if level == 1 and skip_first_title and not first_heading_skipped:
                first_heading_skipped = True
                continue
            style = (
                "Heading 1"
                if level <= 2
                else "Heading 2"
                if level == 3
                else "Heading 3"
            )
            paragraph = document.add_paragraph(style=style)
            _add_word_inline(paragraph, text, style_settings, heading_level=level)
        elif kind == "bullet":
            paragraph = document.add_paragraph(style="List Bullet")
            if level > 0:
                paragraph.paragraph_format.left_indent = Inches(0.75)
            _add_word_inline(paragraph, text, style_settings)
        elif kind == "quote":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.18)
            paragraph.paragraph_format.right_indent = Inches(0.12)
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(5)
            _shade_paragraph(paragraph, LIGHT_FILL)
            _add_left_border(paragraph, ACCENT)
            _add_word_inline(paragraph, text, style_settings)
        elif kind == "table":
            _add_word_table(document, text, style_settings)
        else:
            paragraph = document.add_paragraph()
            _add_word_inline(paragraph, text, style_settings)


def _markdown_blocks(markdown: str):
    lines = markdown.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        if not line.strip():
            index += 1
            continue
        if (
            line.lstrip().startswith("|")
            and index + 1 < len(lines)
            and _is_markdown_table_separator(lines[index + 1])
        ):
            rows = [_table_cells(line)]
            index += 2
            while index < len(lines) and lines[index].lstrip().startswith("|"):
                rows.append(_table_cells(lines[index]))
                index += 1
            width = max(len(row) for row in rows)
            yield "table", [row + [""] * (width - len(row)) for row in rows], 0
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            yield "heading", _strip_leading_punctuation(heading.group(2)), len(heading.group(1))
            index += 1
            continue
        quote = re.match(r"^>\s?(.*)$", line)
        if quote:
            yield "quote", _strip_leading_punctuation(quote.group(1)), 0
            index += 1
            continue
        bullet = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if bullet:
            yield "bullet", _strip_leading_punctuation(bullet.group(2)), 1 if len(bullet.group(1)) >= 2 else 0
            index += 1
            continue
        yield "paragraph", _strip_leading_punctuation(line), 0
        index += 1


def _is_markdown_table_separator(line: str) -> bool:
    cells = _table_cells(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def _table_cells(line: str) -> list[str]:
    return [
        cell.strip()
        if re.fullmatch(r":?-{3,}:?", cell.strip())
        else _strip_leading_punctuation(cell.strip())
        for cell in line.strip().strip("|").split("|")
    ]


def _add_word_table(
    document: Document,
    rows: list[list[str]],
    style_settings: ReportStyleSettings,
) -> None:
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = True
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(2)
            _add_word_inline(paragraph, value, style_settings)
            if row_index == 0:
                _shade_cell(cell, ACCENT)
            for run in paragraph.runs:
                _set_run_font(
                    run,
                    *_word_fonts(style_settings.font_label),
                    max(8.2, _px_to_pt(style_settings.body_size) - 1.5),
                    "FFFFFF" if row_index == 0 else _hex(style_settings.body_color, INK),
                    bold=row_index == 0,
                )
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def _shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


INLINE_PATTERN = re.compile(r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\([^)]+\))")


def _add_word_inline(
    paragraph,
    text: str,
    style_settings: ReportStyleSettings,
    *,
    heading_level: int | None = None,
) -> None:
    latin_font, east_asia_font = _word_fonts(style_settings.font_label)
    body_size = _px_to_pt(style_settings.body_size)
    body_color = _hex(style_settings.body_color, MUTED)
    if heading_level is not None:
        body_size = _px_to_pt(
            style_settings.level_one_size
            if heading_level <= 2
            else style_settings.level_two_size
            if heading_level == 3
            else max(17, style_settings.level_two_size - 3)
        )
        body_color = _hex(style_settings.heading_color, INK)
    cursor = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            _set_run_font(run, latin_font, east_asia_font, body_size, body_color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            _set_run_font(run, latin_font, east_asia_font, body_size, body_color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            _set_run_font(run, "Courier New", east_asia_font, max(9.5, body_size - 1), ACCENT_DARK)
        else:
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            _add_hyperlink(paragraph, label, url)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        _set_run_font(run, latin_font, east_asia_font, body_size, body_color)


def _add_hyperlink(paragraph, label: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), ACCENT)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([color, underline])
    text = OxmlElement("w:t")
    text.text = label
    run.extend([properties, text])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _set_style_font(style, latin: str, east_asia: str, size: float, color: str, bold=False):
    style.font.name = latin
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    style._element.rPr.rFonts.set(qn("w:ascii"), latin)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), latin)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def _set_run_font(run, latin: str, east_asia: str, size: float, color: str, bold=False):
    run.font.name = latin
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.bold = bold
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)


def _add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def _shade_paragraph(paragraph, fill: str) -> None:
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _add_left_border(paragraph, color: str) -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        properties.append(borders)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    borders.append(left)


def _pdf_styles(font_name: str, settings: ReportStyleSettings):
    base = getSampleStyleSheet()
    heading_color = colors.HexColor(f"#{_hex(settings.heading_color, INK)}")
    body_color = colors.HexColor(f"#{_hex(settings.body_color, MUTED)}")
    body_size = _px_to_pt(settings.body_size)
    body_leading = body_size * settings.line_height
    h1_size = _px_to_pt(settings.level_one_size)
    h2_size = _px_to_pt(settings.level_two_size)
    cover_size = _px_to_pt(settings.report_title_size)
    return {
        "body": ParagraphStyle(
            "ReportBody", parent=base["BodyText"], fontName=font_name,
            fontSize=body_size, leading=body_leading, textColor=body_color,
            spaceAfter=max(5, body_size * .5), wordWrap="CJK",
        ),
        "h1": ParagraphStyle(
            "ReportH1", fontName=font_name, fontSize=h1_size,
            leading=h1_size * 1.32, textColor=heading_color, spaceBefore=15,
            spaceAfter=8, keepWithNext=True, wordWrap="CJK",
        ),
        "h2": ParagraphStyle(
            "ReportH2", fontName=font_name, fontSize=h2_size,
            leading=h2_size * 1.4, textColor=heading_color, spaceBefore=11,
            spaceAfter=6, keepWithNext=True, wordWrap="CJK",
        ),
        "bullet": ParagraphStyle(
            "ReportBullet", fontName=font_name, fontSize=body_size,
            leading=body_leading, leftIndent=18, firstLineIndent=-10,
            textColor=body_color, spaceAfter=6, wordWrap="CJK",
        ),
        "quote": ParagraphStyle(
            "ReportQuote", fontName=font_name, fontSize=max(9, body_size - .4),
            leading=max(13.5, body_leading - .4), textColor=body_color,
            wordWrap="CJK",
        ),
        "cover_kicker": ParagraphStyle(
            "CoverKicker", fontName=font_name, fontSize=10, leading=13,
            alignment=TA_CENTER, textColor=colors.HexColor(f"#{ACCENT}"),
            spaceAfter=18,
        ),
        "cover_title": ParagraphStyle(
            "CoverTitle", fontName=font_name, fontSize=cover_size,
            leading=cover_size * 1.28, alignment=TA_CENTER,
            textColor=heading_color, spaceAfter=12, wordWrap="CJK",
        ),
        "cover_subtitle": ParagraphStyle(
            "CoverSubtitle", fontName=font_name, fontSize=12, leading=17,
            alignment=TA_CENTER, textColor=colors.HexColor(f"#{ACCENT_DARK}"),
            spaceAfter=24, wordWrap="CJK",
        ),
        "cover_meta": ParagraphStyle(
            "CoverMeta", fontName=font_name, fontSize=9.5, leading=14,
            alignment=TA_CENTER, textColor=body_color, spaceAfter=5,
            wordWrap="CJK",
        ),
    }


def _pdf_cover(context: ReportExportContext, styles) -> list:
    story = [Spacer(1, 1.45 * inch)]
    story.append(Paragraph("INDUSTRY RESEARCH REPORT", styles["cover_kicker"]))
    story.append(Paragraph(_pdf_inline(context.title), styles["cover_title"]))
    story.append(
        Paragraph(
            _pdf_inline(f"{context.industry}  |  {context.region}  |  {context.time_horizon}"),
            styles["cover_subtitle"],
        )
    )
    story.append(Paragraph(_pdf_inline(context.report_status), styles["cover_meta"]))
    if context.sop_label:
        story.append(
            Paragraph(_pdf_inline(f"Methodology: {context.sop_label}"), styles["cover_meta"])
        )
    story.append(
        Paragraph(context.generated_at.strftime("%Y-%m-%d"), styles["cover_meta"])
    )
    return story


def _markdown_to_pdf(markdown: str, styles, *, skip_first_title: bool) -> list:
    story = []
    first_heading_skipped = False
    for kind, text, level in _markdown_blocks(markdown):
        if kind == "heading":
            if level == 1 and skip_first_title and not first_heading_skipped:
                first_heading_skipped = True
                continue
            story.append(Paragraph(_pdf_inline(text), styles["h1" if level <= 2 else "h2"]))
        elif kind == "bullet":
            prefix = "- " if level == 0 else "  - "
            story.append(Paragraph(_pdf_inline(prefix + text), styles["bullet"]))
        elif kind == "quote":
            table = Table(
                [[Paragraph(_pdf_inline(text), styles["quote"]) ]],
                colWidths=[6.72 * inch],
            )
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(f"#{LIGHT_FILL}")),
                        ("BOX", (0, 0), (-1, -1), 0.4, colors.HexColor(f"#{ACCENT}")),
                        ("LEFTPADDING", (0, 0), (-1, -1), 8),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                        ("TOPPADDING", (0, 0), (-1, -1), 6),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ]
                )
            )
            story.extend([table, Spacer(1, 5)])
        elif kind == "table":
            story.extend([_pdf_table(text, styles), Spacer(1, 7)])
        else:
            story.append(Paragraph(_pdf_inline(text), styles["body"]))
    return story


def _pdf_table(rows: list[list[str]], styles) -> Table:
    column_count = len(rows[0])
    data = []
    for row_index, row in enumerate(rows):
        rendered_row = []
        for cell in row:
            content = _pdf_inline(cell)
            if row_index == 0:
                content = f"<font color='#FFFFFF'><b>{content}</b></font>"
            rendered_row.append(Paragraph(content, styles["body"]))
        data.append(rendered_row)
    table = Table(
        data,
        colWidths=[6.72 * inch / column_count] * column_count,
        repeatRows=1,
    )
    commands = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{ACCENT}")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#C9D4D8")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]
    for row_index in range(2, len(rows), 2):
        commands.append(
            ("BACKGROUND", (0, row_index), (-1, row_index), colors.HexColor("#F7FAFA"))
        )
    table.setStyle(TableStyle(commands))
    return table


def _pdf_inline(text: str) -> str:
    pieces: list[str] = []
    cursor = 0
    for match in INLINE_PATTERN.finditer(text):
        pieces.append(html.escape(text[cursor:match.start()]))
        token = match.group(0)
        if token.startswith("**"):
            pieces.append(f"<b>{html.escape(token[2:-2])}</b>")
        elif token.startswith("`"):
            pieces.append(f"<font color='#{ACCENT_DARK}'>{html.escape(token[1:-1])}</font>")
        else:
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            pieces.append(
                f"<link href='{html.escape(url, quote=True)}' color='#{ACCENT}'>"
                f"{html.escape(label)}</link>"
            )
        cursor = match.end()
    pieces.append(html.escape(text[cursor:]))
    return "".join(pieces)


class _NumberedCanvas(canvas.Canvas):
    def __init__(self, *args, report_title: str, **kwargs):
        super().__init__(*args, **kwargs)
        self.report_title = report_title
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        page_count = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self._draw_page_number(page_count)
            super().showPage()
        super().save()

    def _draw_page_number(self, page_count: int):
        page_number = self._pageNumber
        if page_number == 1:
            return
        self.saveState()
        self.setStrokeColor(colors.HexColor("#D7E0E2"))
        self.setLineWidth(0.4)
        self.line(0.82 * inch, 0.56 * inch, 7.68 * inch, 0.56 * inch)
        self.setFont("Helvetica", 8)
        self.setFillColor(colors.HexColor(f"#{MUTED}"))
        self.drawString(0.82 * inch, 0.38 * inch, "Trident")
        self.drawRightString(7.68 * inch, 0.38 * inch, f"{page_number} / {page_count}")
        self.restoreState()
